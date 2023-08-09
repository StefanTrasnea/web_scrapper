import requests
from bs4 import BeautifulSoup
import pandas as pd
import docx

# URL of the webpage with the tables
url = "https://www.loto.ro/loto-new/newLotoSiteNexioFinalVersion/web/app2.php/jocuri/649_si_noroc/rezultate_extragere.html"

# Send an HTTP GET request
response = requests.get(url)

# Parse the HTML content using BeautifulSoup
soup = BeautifulSoup(response.content, "html.parser")

# Find all table elements on the page
tables = soup.find_all("table")

# Initialize a list to store data from both tables
all_data = []
counter = 0
# Iterate through each table
for table in tables:
    if counter < 2:
        table_data = []
        # Iterate through rows in the table
        for row in table.find_all("tr"):
            row_data = []
            for cell in row.find_all(["th", "td"]):
                row_data.append(cell.get_text(strip=True))
            table_data.append(row_data)

        all_data.append(table_data)
        counter +=1
# Print the scraped data from both tables

doc = docx.Document()

# Define font size
font_size = docx.shared.Pt(16)  # Change the value as needed

# Add tables to the Word document
for table_data in all_data:

    # Add spacing between tables
    doc.add_paragraph("")  # Add an empty paragraph for spacing        

    table = doc.add_table(rows=1, cols=len(table_data[0]))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(table_data[0]):
        hdr_cells[idx].text = header

    for row_data in table_data[1:]:
        row_cells = table.add_row().cells
        for idx, cell_data in enumerate(row_data):
            row_cells[idx].text = cell_data



# Save the Word document
doc.save("scraped_data.docx")

